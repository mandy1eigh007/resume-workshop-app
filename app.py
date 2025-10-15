# app.py — Resume Workshop & Pathways (Seattle Tri-County)
# Streamlit single-file app. No APIs. Browser-only. Python 3.11.
# This build:
# - Event-driven Autofill: runs automatically when uploads/URLs/paste change (button kept for manual retry)
# - Label-aware header parsing (Name:/Phone:/Email:/City:/State:) + regex (email/phone/city-state)
# - Cached file text extraction (st.cache_data) with pypdf → pdfminer.six fallback
# - Stronger company vs City, ST guard; tighter date parsing
# - Crew-forward objective language; banned-terms scrub
# - Skills: auto-populate all three buckets + Quick Add kept
# - Expanded feeder-role seeds (line cook, retail, warehouse, barista, server, janitor, custodian, military, driver, landscaper, security, housekeeper, mover)
# - Certifications normalizer (OSHA-10, Flagger (WA), Forklift, CPR, First Aid, etc.)
# - “Clear Autofill” (resets fields only if they still equal the last parsed value)
# - Instructor Packet: TOC + Sources table + verbatim full-text + optional Roadmap appendix (slice of uploaded *Roadmaps* DOCX for selected trade)
# - Removed obsolete “Step 9” draft section

from __future__ import annotations
import io, os, re, csv, hashlib, datetime
from dataclasses import dataclass, asdict
from typing import List, Dict, Any, Tuple

import streamlit as st
import pandas as pd
from docxtpl import DocxTemplate
from docx import Document as DocxWriter
from docx.shared import Pt, Inches
from pypdf import PdfReader
import requests

# Optional PDF fallback (no error if missing)
try:
    from pdfminer.high_level import extract_text as pdfminer_extract_text
except Exception:
    pdfminer_extract_text = None

st.set_page_config(page_title="Resume Workshop & Pathways", layout="wide")

# ─────────────────────────────────────────────────────────
# Constants / regex
# ─────────────────────────────────────────────────────────
MAX_SUMMARY_CHARS = 450
MAX_SKILLS = 12
MAX_CERTS = 12
MAX_JOBS = 3
MAX_BULLETS_PER_JOB = 4
MAX_SCHOOLS = 2

UNION_BANS = [
    r"\bunion\b", r"\bnon[-\s]?union\b", r"\bibew\b", r"\blocal\s*\d+\b",
    r"\binside\s*wire(man|men)?\b", r"\blow[-\s]?voltage\b", r"\bsound\s+and\s+communication(s)?\b",
    r"\bneca\b", r"\bopen[-\s]?shop\b"
]
BANNED_RE = re.compile("|".join(UNION_BANS), re.I)
FILLER_LEADS = re.compile(r"^\s*(responsible for|duties included|tasked with|in charge of)\s*:?\s*", re.I)
MULTISPACE = re.compile(r"\s+")
EMAIL_RE = re.compile(r"[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}", re.I)
PHONE_RE = re.compile(r"(\+?1[\s\-\.])?\(?\d{3}\)?[\s\-\.]?\d{3}[\s\-\.]?\d{4}")
PHONE_DIGITS = re.compile(r"\D+")
CITY_STATE_RE = re.compile(r"\b([A-Za-z .'-]{2,}),\s*([A-Za-z]{2})\b")
DATE_RANGE_RE = re.compile(
    r"(?P<start>(?:\d{4}|\w{3,9}\s+\d{4}))\s*(?:–|-|to|until|through)\s*(?P<end>(?:Present|Current|\d{4}|\w{3,9}\s+\d{4}))",
    re.I
)
LABEL_RE = re.compile(r"^\s*(name|phone|email|city|state)\s*:\s*(.+)$", re.I)

# ─────────────────────────────────────────────────────────
# Basic cleaners
# ─────────────────────────────────────────────────────────
def strip_banned(text: str) -> str:
    return BANNED_RE.sub("", text or "").strip()

def norm_ws(s: str) -> str:
    s = (s or "").strip()
    return MULTISPACE.sub(" ", s)

def cap_first(s: str) -> str:
    s = norm_ws(s)
    return s[:1].upper()+s[1:] if s else s

def clean_phone(s: str) -> str:
    digits = PHONE_DIGITS.sub("", s or "")
    if len(digits)==11 and digits.startswith("1"): digits = digits[1:]
    if len(digits)==10: return f"({digits[:3]}) {digits[3:6]}-{digits[6:]}"
    return norm_ws(s or "")

def clean_email(s: str) -> str:
    return (s or "").strip().lower()

def clean_bullet(s: str) -> str:
    s = norm_ws(s)
    s = re.sub(r"^[•\-\u2022]+\s*", "", s)
    s = FILLER_LEADS.sub("", s)
    s = re.sub(r"\.+$","", s)
    s = cap_first(s)
    words = s.split()
    return " ".join(words[:24]) if len(words)>24 else s

def split_list(raw: str) -> List[str]:
    if not raw: return []
    parts = [p.strip(" •\t") for p in re.split(r"[,\n;•]+", raw)]
    return [p for p in parts if p]

def parse_dates(raw: str) -> Tuple[str,str]:
    raw = norm_ws(raw)
    m = DATE_RANGE_RE.search(raw)
    if m: return (m.group("start"), m.group("end"))
    if "–" in raw or "-" in raw:
        sep = "–" if "–" in raw else "-"
        bits = [b.strip() for b in raw.split(sep,1)]
        if len(bits)==2: return bits[0], bits[1]
    return (raw,"") if raw else ("","")

# ─────────────────────────────────────────────────────────
# File text extraction & public URL fetch (cached)
# ─────────────────────────────────────────────────────────
class NamedBytesIO(io.BytesIO):
    def __init__(self, data: bytes, name: str):
        super().__init__(data)
        self.name = name

def _drive_direct(url: str) -> str:
    m = re.search(r"/d/([a-zA-Z0-9_-]+)", url) or re.search(r"[?&]id=([a-zA-Z0-9_-]+)", url)
    if m:
        file_id = m.group(1)
        return f"https://drive.google.com/uc?export=download&id={file_id}"
    return url

def fetch_url_to_named_bytes(url: str, fallback_name: str = "downloaded_file") -> NamedBytesIO | None:
    try:
        u = _drive_direct(url.strip())
        r = requests.get(u, timeout=30)
        r.raise_for_status()
        name = (url.split("?")[0].rstrip("/").split("/")[-1] or fallback_name)
        if "." not in name:
            ct = r.headers.get("content-type","").lower()
            if "pdf" in ct: name += ".pdf"
            elif "word" in ct or "officedocument" in ct: name += ".docx"
            else: name += ".txt"
        return NamedBytesIO(r.content, name)
    except Exception:
        return None

def _hash_bytes(b: bytes) -> str:
    return hashlib.sha256(b).hexdigest()

@st.cache_data(show_spinner=False)
def _cached_extract_text(name: str, file_hash: str, kind: str, raw: bytes) -> str:
    # name/kind/file_hash used for cache identity; raw kept for extraction
    if kind == "pdf":
        try:
            reader = PdfReader(io.BytesIO(raw)); chunks=[]
            for p in reader.pages:
                txt = p.extract_text() or ""
                chunks.append(txt)
            text = "\n".join(chunks)
            if text.strip(): return text
        except Exception:
            pass
        if pdfminer_extract_text is not None:
            try:
                text = pdfminer_extract_text(io.BytesIO(raw)) or ""
                return text
            except Exception:
                return ""
        return ""
    elif kind == "docx":
        try:
            from docx import Document as DocxReader
            doc = DocxReader(io.BytesIO(raw))
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception:
            return ""
    else:
        try:
            return raw.decode("utf-8", errors="ignore")
        except Exception:
            return ""

def extract_text_generic(upload) -> str:
    # Works for Streamlit UploadedFile and our NamedBytesIO
    name = getattr(upload, "name", "") or "file"
    lname = name.lower()
    if hasattr(upload, "getvalue"):
        raw = upload.getvalue()
    else:
        raw = upload.read()
    kind = "txt"
    if lname.endswith(".pdf"): kind = "pdf"
    elif lname.endswith(".docx"): kind = "docx"
    h = _hash_bytes(raw)
    return _cached_extract_text(name, h, kind, raw)

# ─────────────────────────────────────────────────────────
# Skills canon & mining
# ─────────────────────────────────────────────────────────
SKILL_CANON = [
    "Problem-solving","Critical thinking","Attention to detail","Time management",
    "Teamwork & collaboration","Adaptability & willingness to learn","Safety awareness",
    "Conflict resolution","Customer service","Leadership","Reading blueprints & specs",
    "Hand & power tools","Materials handling (wood/concrete/metal)","Operating machinery",
    "Trades math & measurement","Regulatory compliance","Physical stamina & dexterity"
]
_SKILL_SYNONYMS = {
    "problem solving":"Problem-solving","problem-solving":"Problem-solving",
    "critical-thinking":"Critical thinking","attention to details":"Attention to detail",
    "time-management":"Time management","teamwork":"Teamwork & collaboration",
    "collaboration":"Teamwork & collaboration","adaptability":"Adaptability & willingness to learn",
    "willingness to learn":"Adaptability & willingness to learn","safety":"Safety awareness",
    "customer service skills":"Customer service","leadership skills":"Leadership",
    "blueprints":"Reading blueprints & specs","tools":"Hand & power tools",
    "machinery":"Operating machinery","math":"Trades math & measurement",
    "measurements":"Trades math & measurement","compliance":"Regulatory compliance",
    "stamina":"Physical stamina & dexterity","forklift":"Operating machinery",
}
TRANSFERABLE_KEYWORDS = {
    "problem": "Problem-solving", "solve": "Problem-solving", "troubleshoot": "Problem-solving",
    "analyz": "Critical thinking", "priorit": "Time management", "deadline": "Time management",
    "detail": "Attention to detail", "team": "Teamwork & collaboration", "collabor": "Teamwork & collaboration",
    "adapt": "Adaptability & willingness to learn", "learn": "Adaptability & willingness to learn",
    "safety": "Safety awareness", "osha": "Safety awareness", "customer": "Customer service",
    "lead": "Leadership", "blueprint": "Reading blueprints & specs", "spec": "Reading blueprints & specs",
    "tool": "Hand & power tools", "drill": "Hand & power tools", "saw": "Hand & power tools",
    "forklift": "Operating machinery", "material": "Materials handling (wood/concrete/metal)",
    "machin": "Operating machinery", "math": "Trades math & measurement", "measure": "Trades math & measurement",
    "code": "Regulatory compliance", "permit": "Regulatory compliance", "compliance": "Regulatory compliance",
    "stamina": "Physical stamina & dexterity", "lift": "Physical stamina & dexterity",
}

def normalize_skill_label(s: str) -> str:
    base = (s or "").strip()
    key = re.sub(r"\s+"," ",base.lower())
    mapped = _SKILL_SYNONYMS.get(key)
    if mapped: return mapped
    return re.sub(r"\s+"," ",base).strip().title()

def suggest_transferable_skills_from_text(text: str) -> List[str]:
    hits = {}
    low = (text or "").lower()
    for kw, skill in TRANSFERABLE_KEYWORDS.items():
        if kw in low:
            hits[skill] = hits.get(skill, 0) + 1
    ordered = [s for s,_ in sorted(hits.items(), key=lambda kv: -kv[1])]
    canon_order = [s for s in SKILL_CANON if s in ordered]
    return canon_order[:8]

def categorize_skills(skills: List[str]) -> Dict[str, List[str]]:
    out = {"Transferable": [], "Job-Specific": [], "Self-Management": []}
    seen=set()
    for s in skills:
        lab = normalize_skill_label(s)
        if not lab: continue
        if lab.lower() in seen: continue
        seen.add(lab.lower())
        if lab in {"Reading blueprints & specs","Hand & power tools","Operating machinery","Materials handling (wood/concrete/metal)","Trades math & measurement","Regulatory compliance","Safety awareness"}:
            cat="Job-Specific"
        elif lab in {"Leadership","Adaptability & willingness to learn","Physical stamina & dexterity"}:
            cat="Self-Management"
        else:
            cat="Transferable"
        out[cat].append(lab)
    return out

# ─────────────────────────────────────────────────────────
# Trade taxonomy
# ─────────────────────────────────────────────────────────
TRADE_TAXONOMY = [
    "Boilermaker (Local 104, 502)",
    "Bricklayer / BAC Allied (Brick/Tile/Terrazzo/Marble/PCC)",
    "Carpenter (General)",
    "Carpenter – Interior Systems",
    "Millwright",
    "Pile Driver",
    "Cement Mason (OPCMIA 528)",
    "Drywall Finisher (IUPAT)",
    "Electrician – Inside (01)",
    "Electrician – Limited Energy (06)",
    "Electrician – Residential (02)",
    "Elevator Constructor (IUEC/NEIEP)",
    "Floor Layer (IUPAT)",
    "Glazier (IUPAT 188)",
    "Heat & Frost Insulator (Local 7)",
    "Ironworker (Local 86)",
    "Laborer (LIUNA 242/252/292)",
    "Operating Engineer (IUOE 302/612)",
    "Painter (IUPAT DC5)",
    "Plasterer (OPCMIA 528)",
    "Plumber / Steamfitter / HVAC-R (UA 32 / UA 26)",
    "Roofer (Local 54/153)",
    "Sheet Metal (SMART 66)",
    "Sprinkler Fitter (UA 699)",
    "High Voltage – Outside Lineman (NW Line JATC)",
    "Power Line Clearance Tree Trimmer (NW Line JATC)",
]

# ─────────────────────────────────────────────────────────
# Role→construction seed bullets
# ─────────────────────────────────────────────────────────
ROLE_TO_CONSTR_BULLETS = {
    "line cook": [
        "Worked safely around hot equipment and sharp tools",
        "Kept stations clean and organized; followed prep lists",
        "Handled deliveries and rotated stock; maintained clear walkways",
        "Stayed on pace to meet production during rushes",
    ],
    "retail": [
        "Kept inventory organized and aisles clear for safe flow",
        "Communicated with customers and team under time pressure",
        "Handled cash counts and hand-offs accurately",
    ],
    "warehouse": [
        "Staged materials, verified counts, and kept aisles clear",
        "Operated pallet jacks/hand trucks under PPE rules",
    ],
    "barista": [
        "Followed recipes and equipment safety steps precisely",
        "Kept stations stocked and organized during rushes",
    ],
    "server": [
        "Managed multiple tasks with tight timing while supporting team flow",
        "Maintained a clean, safe work area under pressure",
    ],
    "janitor": [
        "Used chemicals/equipment per safety guidance; kept areas hazard-free",
        "Completed checklists and documented work reliably",
    ],
    "custodian": [
        "Set up spaces, moved materials, and followed safety procedures",
    ],
    "military": [
        "Followed procedures, PPE, and safety briefings precisely",
        "Worked in teams with accountability and time standards",
    ],
    "driver": [
        "Loaded/unloaded and secured materials; verified counts",
        "Maintained safe backing/spotting practices on tight sites",
    ],
    "landscaper": [
        "Operated hand tools safely; maintained clean work zones",
        "Staged materials and debris; followed site directions",
    ],
    "security": [
        "Monitored hazards and enforced basic safety practices",
        "Communicated clearly with teams and visitors on-site",
    ],
    "housekeeper": [
        "Kept work areas hazard-free; followed chemical/PPE guidelines",
        "Worked on schedule with checklists and quality checks",
    ],
    "mover": [
        "Lifted and carried materials with safe techniques",
        "Coordinated team moves and protected finished surfaces",
    ],
}

# ─────────────────────────────────────────────────────────
# Parsing helpers (header/jobs/edu/certs)
# ─────────────────────────────────────────────────────────
def parse_header(text: str) -> Dict[str,str]:
    name = ""; email = ""; phone = ""; city = ""; state = ""
    # Label-aware pass (handles templates like ANEW)
    for l in (text or "").splitlines()[:50]:
        m = LABEL_RE.match(l)
        if not m: continue
        key, val = m.group(1).lower(), m.group(2).strip()
        if key=="name" and not name: name = val
        elif key=="phone" and not phone: phone = val
        elif key=="email" and not email: email = val
        elif key=="city" and not city: city = val
        elif key=="state" and not state: state = val
    # Regex pass
    if not email:
        m = EMAIL_RE.search(text or "");  email = m.group(0) if m else ""
    if not phone:
        m = PHONE_RE.search(text or "");  phone = m.group(0) if m else ""
    if not (city and state):
        mcs = CITY_STATE_RE.search(text or "")
        if mcs:
            city, state = mcs.group(1), mcs.group(2).upper()
    # Try name from top contact lines
    if not name:
        top = "\n".join([l.strip() for l in (text or "").splitlines()[:12] if l.strip()])
        sep_line = re.sub(r"[•·–—\-•]+", "|", top)
        for candidate in sep_line.split("\n"):
            if "@" in candidate or re.search(r"\d{3}.*\d{4}", candidate):
                parts = [p.strip() for p in candidate.split("|") if p.strip()]
                for p in parts:
                    if EMAIL_RE.search(p) or PHONE_RE.search(p): continue
                    if CITY_STATE_RE.search(p): continue
                    if 2 <= len(p.split()) <= 4 and not re.search(r"\d", p):
                        caps = sum(w[:1].isupper() for w in p.split())
                        if caps >= 2:
                            name = p
                            break
                if name: break
    return {"Name": cap_first(name), "Email": clean_email(email), "Phone": clean_phone(phone),
            "City": cap_first(city), "State": (state or "").strip().upper()}

def _safe_company_token(token: str) -> bool:
    if CITY_STATE_RE.fullmatch(token):
        return False
    # Likely company if not purely city/state and has letters/spaces/&/'/-
    return bool(re.fullmatch(r"[A-Za-z0-9 .'\-&]{2,}", token))

def parse_jobs(text: str) -> List[Dict[str,Any]]:
    out=[]
    lines = [l.rstrip() for l in (text or "").splitlines()]
    i=0
    while i < len(lines) and len(out) < MAX_JOBS:
        head = lines[i].strip()
        if not head:
            i+=1; continue
        if re.match(r"^\s*(summary|objective|skills|certifications|education)\s*$", head, re.I):
            i+=1; continue

        window = " ".join(lines[i:i+3])
        parts = re.split(r"\s*\|\s*|\s{2,}| — | – ", head)
        role=""; company=""; cityst=""; dates=""

        mdate = DATE_RANGE_RE.search(window)
        if mdate:
            dates = f"{mdate.group('start')} – {mdate.group('end')}"
        mcity = CITY_STATE_RE.search(window)
        if mcity:
            cityst = f"{mcity.group(1)}, {mcity.group(2).upper()}"

        if len(parts) >= 2:
            cand_role = parts[0].strip()
            cand_co   = parts[1].strip()
            if _safe_company_token(cand_co):
                role, company = cand_role, cand_co
            else:
                if _safe_company_token(cand_role) and not _safe_company_token(cand_co):
                    company, role = cand_role, cand_co
                else:
                    role = cand_role
        else:
            role = head

        bullets=[]
        j=i+1
        while j < len(lines):
            ln = lines[j].strip()
            if not ln:
                if bullets: break
                j+=1; continue
            if re.match(r"^\s*(summary|objective|skills|certifications|education)\s*$", ln, re.I):
                break
            if re.match(r"^[•\-\u2022]\s+", ln) or lines[j].startswith("\t"):
                bullets.append(clean_bullet(re.sub(r"^[•\-\u2022\-]+\s*", "", ln)))
                if len(bullets) >= MAX_BULLETS_PER_JOB:
                    break
            else:
                if DATE_RANGE_RE.search(ln) or CITY_STATE_RE.search(ln):
                    pass
                elif len(ln.split()) <= 4 and bullets:
                    break
            j+=1

        if company and CITY_STATE_RE.fullmatch(company):
            company = ""

        if any([company, role, cityst, dates, bullets]):
            out.append({
                "company": cap_first(company),
                "role": cap_first(role),
                "city": cap_first(cityst),
                "start": parse_dates(dates)[0] if dates else "",
                "end": parse_dates(dates)[1] if dates else "",
                "bullets": bullets
            })
        i = max(j, i+1)
    return out[:MAX_JOBS]

def parse_education(text: str) -> List[Dict[str,str]]:
    out=[]
    lines = [l.strip() for l in (text or "").splitlines()]
    i=0
    while i < len(lines) and len(out) < MAX_SCHOOLS:
        l = lines[i]
        if re.search(r"(high school|ged|college|university|program|certificate|diploma)", l, re.I):
            school = cap_first(l)
            cred=""; year=""; details=""
            for la in lines[i+1:i+5]:
                if re.search(r"\b(20\d{2}|19\d{2})\b", la):
                    year = la.strip()
                mcs = CITY_STATE_RE.search(la)
                if mcs:
                    details = f"{mcs.group(1)}, {mcs.group(2).upper()}"
                if any(x in la.lower() for x in ["diploma","degree","certificate","ged","program"]) and not cred:
                    cred = cap_first(la.strip())
            out.append({"school": school, "credential": cred, "year": year, "details": details})
        i+=1
    return out[:MAX_SCHOOLS]

CERT_NORMALIZE = {
    "osha10":"OSHA-10", "osha 10":"OSHA-10", "osha-10":"OSHA-10", "osha":"OSHA-10",
    "fork lift":"Forklift", "fork-lift":"Forklift", "forklift":"Forklift",
    "flagger (wa)":"Flagger (WA)", "flagger wa":"Flagger (WA)", "flagger":"Flagger",
    "cpr":"CPR", "first aid":"First Aid", "first-aid":"First Aid", "cpr/first aid":"CPR, First Aid",
    "aerial lift":"Aerial Lift", "hazwoper":"HAZWOPER", "twic":"TWIC",
    "confined space":"Confined Space", "traffic control":"Traffic Control",
    "nccer":"NCCER", "ppe":"PPE"
}
CERT_KEYWORDS = list(set(list(CERT_NORMALIZE.keys()) + [
    "osha", "forklift", "flagger", "cpr", "first aid", "hazwoper",
    "twic", "nccer", "confined space", "ppe", "aerial lift", "traffic control"
]))
def _normalize_cert_token(tok: str) -> List[str]:
    t = tok.strip().lower()
    out = CERT_NORMALIZE.get(t)
    if out:
        return [c.strip() for c in out.split(",")]
    # smart splits like "CPR / First Aid"
    t2 = re.split(r"[\/,;]+", t)
    if len(t2) > 1:
        res=[]
        for piece in t2:
            res += _normalize_cert_token(piece)
        return res
    # default title
    return [tok.strip().title()]

def parse_certs(text: str) -> List[str]:
    found=set()
    low = (text or "").lower()
    for k in CERT_KEYWORDS:
        if k in low:
            found.update(_normalize_cert_token(k))
    # explicit scan of lines
    for line in (text or "").splitlines():
        for k in CERT_KEYWORDS:
            if re.search(rf"\b{k}\b", line, re.I):
                for c in _normalize_cert_token(k):
                    found.add(c)
    # collapse known combos
    out = sorted(found)
    # merge duplicate forms like "CPR" and "Cpr"
    uniq=[]
    seen=set()
    for c in out:
        key=c.lower()
        if key in seen: continue
        seen.add(key); uniq.append(c)
    # ensure OSHA-10 preferred
    uniq=[ "OSHA-10" if u.lower() in {"osha","osha 10","osha-10","osha10"} else u for u in uniq ]
    return uniq[:MAX_CERTS]

def parse_skills_from_text(text: str) -> Dict[str, List[str]]:
    base = suggest_transferable_skills_from_text(text)
    cat = categorize_skills(base)
    return cat

# ─────────────────────────────────────────────────────────
# Objective generator (crew-forward)
# ─────────────────────────────────────────────────────────
def build_objective(trade: str, pitch: str, skills_cat: Dict[str,List[str]]) -> str:
    top = (skills_cat.get("Job-Specific", []) + skills_cat.get("Transferable", []))[:3]
    picks = ", ".join(top) if top else "safety, teamwork, and reliable production"
    p = norm_ws(pitch or "")
    core = f"Ready to get on a crew in {trade}—bringing {picks}. "
    if p:
        core += f"{p} "
    core += "Show up, work safe, learn fast, and help the crew hit targets."
    core = strip_banned(core)[:MAX_SUMMARY_CHARS]
    return core

# ─────────────────────────────────────────────────────────
# Data classes + resume rendering
# ─────────────────────────────────────────────────────────
@dataclass
class Job:
    company:str=""; role:str=""; city:str=""; start:str=""; end:str=""; bullets:List[str]=None
    def trim(self,k:int):
        bs=[clean_bullet(b) for b in (self.bullets or []) if str(b).strip()]
        self.bullets = bs[:k]

@dataclass
class School:
    school:str=""; credential:str=""; year:str=""; details:str=""

def build_resume_context(form: Dict[str,Any], trade_label: str) -> Dict[str,Any]:
    Name=cap_first(form["Name"]); City=cap_first(form["City"]); State=(form["State"] or "").strip().upper()
    phone=clean_phone(form["Phone"]); email=clean_email(form["Email"])

    if form.get("Objective_Final","").strip():
        summary = strip_banned(norm_ws(form["Objective_Final"]))[:MAX_SUMMARY_CHARS]
    else:
        skills_cat = categorize_skills(
            split_list(form.get("Skills_Transferable","")) +
            split_list(form.get("Skills_JobSpecific","")) +
            split_list(form.get("Skills_SelfManagement",""))
        )
        summary = build_objective(trade_label, form.get("Pitch",""), skills_cat)

    # Skills
    skills_all=[]
    for raw in (form["Skills_Transferable"], form["Skills_JobSpecific"], form["Skills_SelfManagement"]):
        skills_all += split_list(raw)
    seen=set(); skills=[]
    for s in skills_all:
        lab=normalize_skill_label(norm_ws(s))
        if lab and lab.lower() not in seen:
            seen.add(lab.lower()); skills.append(lab)
    skills = skills[:MAX_SKILLS]

    certs = [norm_ws(c) for c in split_list(form["Certifications"] )][:MAX_CERTS]

    # Jobs
    jobs=[]
    for idx in range(1, MAX_JOBS+1):
        company=form.get(f"Job{idx}_Company",""); cityst=form.get(f"Job{idx}_CityState","")
        dates=form.get(f"Job{idx}_Dates",""); title=form.get(f"Job{idx}_Title",""); duties=form.get(f"Job{idx}_Duties","")
        if not any([company, title, duties, dates, cityst]): continue
        s,e = parse_dates(dates)
        raw_b = [b for b in (duties or "").splitlines() if b.strip()]
        j = Job(company=cap_first(company), role=cap_first(title), city=cap_first(cityst),
                start=norm_ws(s), end=norm_ws(e), bullets=raw_b)
        j.trim(MAX_BULLETS_PER_JOB); jobs.append(j)
    jobs = jobs[:MAX_JOBS]

    # Schools
    schools=[]
    for idx in range(1, MAX_SCHOOLS+1):
        sch=form.get(f"Edu{idx}_School",""); cs=form.get(f"Edu{idx}_CityState",""); d=form.get(f"Edu{idx}_Dates",""); cr=form.get(f"Edu{idx}_Credential","")
        if not any([sch,cr,d,cs]): continue
        schools.append(School(school=cap_first(sch), credential=cap_first(cr), year=norm_ws(d), details=cap_first(cs) if cs else ""))

    other_work = norm_ws(form.get("Other_Work",""))
    volunteer  = norm_ws(form.get("Volunteer",""))
    tail=[]
    if other_work: tail.append(f"Other work: {other_work}")
    if volunteer:  tail.append(f"Volunteer: {volunteer}")
    if tail:
        add="  •  ".join(tail)
        summary = (summary + " " + add).strip()[:MAX_SUMMARY_CHARS]

    return {
        "Name": Name, "City": City, "State": State,
        "phone": phone, "email": email, "summary": summary,
        "skills": skills,
        "certs": certs,
        "jobs": [asdict(j) for j in jobs if any([j.company,j.role,j.bullets])],
        "schools": [asdict(s) for s in schools if any([s.school,s.credential,s.year,s.details])],
        "trade_label": trade_label,
    }

def render_docx_with_template(template_bytes: bytes, context: Dict[str,Any]) -> bytes:
    tpl = DocxTemplate(io.BytesIO(template_bytes))
    tpl.render(context)
    out = io.BytesIO(); tpl.save(out); out.seek(0)
    return out.getvalue()

def _split_highlights(raw: str) -> List[str]:
    return split_list(raw)

def build_cover_letter_docx(data: Dict[str,str]) -> bytes:
    role = strip_banned(data.get("role",""))
    company = strip_banned(data.get("company",""))
    body_strength = strip_banned(data.get("strength",""))
    trade_label = strip_banned(data.get("trade_label",""))
    app_type = (data.get("application_type","Apprenticeship") or "Apprenticeship").strip()

    doc = DocxWriter()
    styles = doc.styles['Normal']; styles.font.name = 'Calibri'; styles.font.size = Pt(11)

    doc.add_paragraph(f"{data.get('name','')}")
    doc.add_paragraph(f"{data.get('city','')}, {data.get('state','')}")
    contact = ", ".join([x for x in [data.get('phone',''), data.get('email','')] if x])
    if contact: doc.add_paragraph(contact)
    doc.add_paragraph("")

    today = datetime.date.today().strftime("%B %d, %Y")
    doc.add_paragraph(today)
    if company: doc.add_paragraph(company)
    if data.get("location"): doc.add_paragraph(data["location"])
    doc.add_paragraph("")
    doc.add_paragraph("Dear Hiring Committee,")

    p1 = doc.add_paragraph()
    p1.add_run(
        f"I’m applying for a {role} {('apprenticeship' if app_type=='Apprenticeship' else 'position')} "
        f"in the {trade_label} scope. I bring reliability, safety awareness, and hands-on readiness."
    )
    p2 = doc.add_paragraph()
    p2.add_run(
        "My background includes tool proficiency, teamwork under real schedules, and a commitment to quality and safe production."
    )
    hi = _split_highlights(body_strength)
    if hi:
        doc.add_paragraph("Highlights:")
        for line in hi:
            doc.add_paragraph(f"• {line}")

    doc.add_paragraph("")
    doc.add_paragraph("Thank you for your consideration. I’m ready to support your crew and learn the trade the right way.")
    doc.add_paragraph("")
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph(data.get("name",""))

    bio = io.BytesIO(); doc.save(bio); bio.seek(0)
    return bio.getvalue()

# ────────── Instructor Packet: TOC + Sources + Full Text + Roadmap slice ──────────
def _add_toc(doc: DocxWriter, entries: List[str]):
    doc.add_heading("Table of Contents", level=1)
    for i, e in enumerate(entries, 1):
        doc.add_paragraph(f"{i}. {e}")

def _add_sources_table(doc: DocxWriter, sources: List[Any]):
    doc.add_heading("Sources Imported", level=1)
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "File"
    hdr[1].text = "Type"
    hdr[2].text = "Imported"
    now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
    for s in sources:
        name = getattr(s, "name", "file")
        typ = os.path.splitext(name)[1].lower().lstrip(".") or "txt"
        row = table.add_row().cells
        row[0].text = name
        row[1].text = typ
        row[2].text = now

def _roadmap_slice_from_docx(doc_bytes: bytes, trade_label: str) -> List[str]:
    """Return paragraph list for the selected trade from a Roadmaps DOCX."""
    try:
        doc = DocxWriter(io.BytesIO(doc_bytes))
        paras = [p.text.strip() for p in doc.paragraphs]
    except Exception:
        return []
    # Find a section where the trade label appears; gather until next all-caps heading or blank gap
    start_idx = -1
    for idx, t in enumerate(paras):
        if not t: continue
        if trade_label.lower() in t.lower():
            start_idx = idx
            break
    if start_idx < 0:
        return []
    block = []
    for t in paras[start_idx:]:
        if t and t.isupper() and len(t.split()) <= 12 and t.lower() != trade_label.lower():
            # likely a new trade heading → stop
            break
        block.append(t)
    # Clean leading heading duplication
    return [x for x in block if x is not None]

def build_pathway_packet_docx(student: Dict[str,str], trade_label: str, app_type: str, sources: List[Any], reflections: Dict[str,str]) -> bytes:
    doc = DocxWriter()
    styles = doc.styles['Normal']; styles.font.name = 'Calibri'; styles.font.size = Pt(11)

    toc_entries = ["Workshop Reflections", "Full Text of Uploaded/Imported Files"]
    doc.add_heading("Instructor Pathway Packet", level=0)
    meta = f"Student: {student.get('name','')} | Target: {trade_label} | Application type: {app_type}"
    doc.add_paragraph(meta); doc.add_paragraph("")

    # If any source looks like a Roadmap docx, we’ll append a focused slice later
    roadmap_slice: List[str] = []
    for upl in sources or []:
        nm = getattr(upl, "name", "").lower()
        if nm.endswith(".docx") and "roadmap" in nm:
            try:
                # get bytes
                raw = upl.getvalue() if hasattr(upl, "getvalue") else upl.read()
                roadmap_slice = _roadmap_slice_from_docx(raw, trade_label)
                break
            except Exception:
                pass
    if roadmap_slice:
        toc_entries.append("Trade Roadmap (Relevant Excerpt)")

    _add_toc(doc, toc_entries)

    # Sources table
    _add_sources_table(doc, sources)

    doc.add_page_break()
    doc.add_heading("Workshop Reflections", level=1)
    for k,v in reflections.items():
        doc.add_paragraph(k+":")
        for line in (v or "").splitlines():
            doc.add_paragraph(line)

    # Full text of files
    doc.add_page_break()
    doc.add_heading("Full Text of Uploaded/Imported Files", level=1)
    for upl in sources or []:
        doc.add_page_break()
        doc.add_heading(getattr(upl,"name","(file)"), level=2)
        text = extract_text_generic(upl)
        if text.strip():
            for line in text.splitlines():
                doc.add_paragraph(line)
        else:
            doc.add_paragraph("Couldn’t extract text. Tip: upload as text-based PDF or DOCX, not scans.")

    # Roadmap slice (optional)
    if roadmap_slice:
        doc.add_page_break()
        doc.add_heading("Trade Roadmap (Relevant Excerpt)", level=1)
        for line in roadmap_slice:
            doc.add_paragraph(line)

    out = io.BytesIO(); doc.save(out); out.seek(0)
    return out.getvalue()

# ─────────────────────────────────────────────────────────
# Sidebar — template + extra docs
# ─────────────────────────────────────────────────────────
with st.sidebar:
    st.header("Templates & Docs")

    tpl_bytes=None
    if os.path.exists("resume_app_template.docx"):
        with open("resume_app_template.docx","rb") as f: tpl_bytes=f.read()
    upl_tpl = st.file_uploader("Upload RESUME DOCX template (optional)", type=["docx"])
    if upl_tpl is not None:
        tpl_bytes = upl_tpl.read()

    st.markdown("---")
    st.caption("Upload additional instructor/pathway docs (PDF/DOCX/TXT). These get embedded (full text) in the Instructor Packet.")
    pathway_uploads = st.file_uploader("Upload pathway documents", type=["pdf","docx","txt"], accept_multiple_files=True)

# ─────────────────────────────────────────────────────────
# Main — Step 0: Intake (uploads/URLs/paste)
# ─────────────────────────────────────────────────────────
st.title("Student Packet: Resume Workshop")

st.subheader("0) Bring Your Stuff (we’ll mine it)")
c0a, c0b = st.columns(2)
with c0a:
    prev_resume_files = st.file_uploader(
        "Previous resume (PDF/DOCX/TXT)", type=["pdf","docx","txt"], accept_multiple_files=True
    )
with c0b:
    jd_files = st.file_uploader(
        "Job descriptions / postings (PDF/DOCX/TXT)", type=["pdf","docx","txt"], accept_multiple_files=True
    )

st.markdown("**Or import by URL (public links only: Google Drive/GCS):**")
url_list = st.text_area("One URL per line", "")
url_fetches = []
if url_list.strip():
    for i, raw in enumerate(url_list.splitlines(), start=1):
        u = raw.strip()
        if not u:
            continue
        nb = fetch_url_to_named_bytes(u, fallback_name=f"url_{i}")
        if nb is not None:
            url_fetches.append(nb)
        else:
            st.warning(f"Could not fetch: {u} (is it public?)")

st.markdown("**Or paste text directly:**")
paste_box = st.text_area("Paste any resume or job description text here", "")

def extract_multi(files) -> str:
    txt = ""
    for f in files or []:
        txt += "\n" + extract_text_generic(f)
    return txt

prev_resume_text = extract_multi(prev_resume_files)
jd_text_files = extract_multi(jd_files) + extract_multi(url_fetches)
jd_text_paste = paste_box or ""
combined_text = "\n".join([prev_resume_text, jd_text_files, jd_text_paste]).strip()

if combined_text:
    st.caption(f"Loaded text from {len(prev_resume_files or [])} resume file(s), "
               f"{len(jd_files or []) + len(url_fetches)} JD file(s)/URL(s), "
               f"and {'pasted text' if jd_text_paste else 'no pasted text'}.")
    preview = combined_text[:1000].replace("\n"," ")
    st.info(f"Preview: {preview}…")

# ─────────────────────────────────────────────────────────
# Autofill — automatic + manual; session state + “Clear Autofill”
# ─────────────────────────────────────────────────────────
if "autofilled" not in st.session_state:
    st.session_state["autofilled"] = False
if "autofill_values" not in st.session_state:
    st.session_state["autofill_values"] = {}  # key -> last parsed value

AUTO_KEYS = [
    "Name","Phone","Email","City","State",
    "Skills_Transferable","Skills_JobSpecific","Skills_SelfManagement",
    "Certifications",
    "Job1_Company","Job1_CityState","Job1_Dates","Job1_Title","Job1_Duties",
    "Job2_Company","Job2_CityState","Job2_Dates","Job2_Title","Job2_Duties",
    "Job3_Company","Job3_CityState","Job3_Dates","Job3_Title","Job3_Duties",
    "Edu1_School","Edu1_CityState","Edu1_Dates","Edu1_Credential",
    "Edu2_School","Edu2_CityState","Edu2_Dates","Edu2_Credential",
    "Objective_Final"
]

def set_if_empty(key: str, val: str):
    if key not in st.session_state or not str(st.session_state.get(key,"")).strip():
        st.session_state[key] = val
        st.session_state["autofill_values"][key] = val  # track provenance

def _apply_role_seed_if_needed(idx: int, role: str, current_bullets: str):
    role_label = (role or "").lower()
    seeds=[]
    for key, arr in ROLE_TO_CONSTR_BULLETS.items():
        if key in role_label:
            seeds = arr[:MAX_BULLETS_PER_JOB]
            break
    if (current_bullets or "").strip():
        return current_bullets
    return "\n".join(seeds) if seeds else current_bullets

def autofill_from_text(text: str, trade_for_objective: str) -> Dict[str, Any]:
    parsed: Dict[str, Any] = {"header":{}, "jobs":[], "schools":[], "certs":[], "skills_cat":{}}

    hdr = parse_header(text)
    parsed["header"] = hdr
    for k,v in {"Name":"Name","Phone":"Phone","Email":"Email","City":"City","State":"State"}.items():
        set_if_empty(v, hdr.get(k,""))

    jobs = parse_jobs(text)
    parsed["jobs"] = jobs
    for idx in range(1, MAX_JOBS+1):
        j = jobs[idx-1] if idx-1 < len(jobs) else {}
        set_if_empty(f"Job{idx}_Company", j.get("company",""))
        set_if_empty(f"Job{idx}_CityState", j.get("city",""))
        dates = " – ".join([x for x in [j.get("start",""), j.get("end","")] if x]).strip(" –")
        set_if_empty(f"Job{idx}_Dates", dates)
        set_if_empty(f"Job{idx}_Title", j.get("role",""))
        # role→seed
        new_b = _apply_role_seed_if_needed(idx, j.get("role",""), st.session_state.get(f"Job{idx}_Duties",""))
        if new_b and not st.session_state.get(f"Job{idx}_Duties","").strip():
            st.session_state[f"Job{idx}_Duties"] = new_b
            st.session_state["autofill_values"][f"Job{idx}_Duties"] = new_b

    schools = parse_education(text)
    parsed["schools"] = schools
    for idx in range(1, MAX_SCHOOLS+1):
        s = schools[idx-1] if idx-1 < len(schools) else {}
        set_if_empty(f"Edu{idx}_School", s.get("school",""))
        set_if_empty(f"Edu{idx}_CityState", s.get("details",""))
        set_if_empty(f"Edu{idx}_Dates", s.get("year",""))
        set_if_empty(f"Edu{idx}_Credential", s.get("credential",""))

    certs = parse_certs(text)
    parsed["certs"] = certs
    if certs:
        val = ", ".join(sorted(set(certs)))
        set_if_empty("Certifications", val)

    sk = parse_skills_from_text(text)
    parsed["skills_cat"] = sk
    # default seeds if nothing detected
    if not any(sk.values()):
        sk = categorize_skills(["Safety awareness","Hand & power tools","Teamwork & collaboration","Time management"])
    if sk.get("Transferable"): set_if_empty("Skills_Transferable", ", ".join(sk["Transferable"]))
    if sk.get("Job-Specific"): set_if_empty("Skills_JobSpecific", ", ".join(sk["Job-Specific"]))
    if sk.get("Self-Management"): set_if_empty("Skills_SelfManagement", ", ".join(sk["Self-Management"]))

    obj = build_objective(trade_for_objective, st.session_state.get("Pitch",""), sk)
    set_if_empty("Objective_Final", obj)

    return parsed

# Auto-run autofill once when text changes
def _content_fingerprint() -> str:
    parts = [
        prev_resume_text[:5000],
        jd_text_files[:5000],
        jd_text_paste[:5000],
        str([getattr(x, "name", "") for x in (prev_resume_files or [])]),
        str([getattr(x, "name", "") for x in (jd_files or [])]),
        str([getattr(x, "name", "") for x in (url_fetches or [])]),
    ]
    return hashlib.md5("||".join(parts).encode("utf-8", errors="ignore")).hexdigest()

if "last_fp" not in st.session_state:
    st.session_state["last_fp"] = ""

parsed_snapshot = None
current_fp = _content_fingerprint()
if combined_text and st.session_state["last_fp"] != current_fp:
    parsed_snapshot = autofill_from_text(combined_text, st.session_state.get("SelectedTrade", "Electrician – Inside (01)"))
    st.session_state["autofilled"] = True
    st.session_state["last_fp"] = current_fp
    st.success("Autofill ran from your uploads/URLs/paste.")

# Manual re-run
cauto1, cauto2, cauto3 = st.columns([1,1,2])
with cauto1:
    if st.button("Re-run Autofill", type="secondary", disabled=(not combined_text)):
        parsed_snapshot = autofill_from_text(combined_text, st.session_state.get("SelectedTrade", "Electrician – Inside (01)"))
        st.session_state["autofilled"] = True
        st.success("Autofill re-ran. Fields updated where empty.")
with cauto2:
    if st.button("Clear Autofill", type="secondary"):
        # Only clear keys whose current value equals the last autofill value
        auto_vals = st.session_state.get("autofill_values", {})
        cleared = []
        for k, v in list(auto_vals.items()):
            if st.session_state.get(k, None) == v:
                st.session_state[k] = ""
                cleared.append(k)
                del st.session_state["autofill_values"][k]
        if cleared:
            st.warning(f"Cleared autofilled fields: {', '.join(cleared)}")
        else:
            st.info("Nothing to clear (fields were edited or not autofilled).")

with st.expander("Autofill Debug (what the parser captured)"):
    if parsed_snapshot:
        st.write(parsed_snapshot)
    else:
        st.caption("No new parse yet in this session, or fields were already filled.")
if st.session_state.get("autofilled"):
    st.info("Autofill status: ON — fields were pre-filled from your uploaded content.")

# ─────────────────────────────────────────────────────────
# Workshop UI (Steps renumbered; Step 9 removed)
# ─────────────────────────────────────────────────────────
st.subheader("What is the Difference Between a Construction Facing Resume and a Traditional Resume?")
st.markdown("""
**Construction Facing Resume**  
• **Purpose:** Getting into a trade, apprenticeship, or construction company.  
• **Focus:** Hands-on skills (tools, materials), certs (OSHA-10, Flagger, Forklift), physical abilities, build projects, and jobsite language.  
• **Experience:** Translate non-construction roles into site value (teamwork, time, safety).
""")
wk_q1 = st.text_area("Write three things you will include on a construction-facing resume that you wouldn’t on a traditional resume:", height=120)

st.subheader("1. Why a Resume Matters in Construction")
st.write("It’s your first proof you can show up safe, use tools, learn fast, and support a crew.")

st.subheader("2. Your Header (Contact Information)")
c1, c2 = st.columns(2)
with c1:
    Name = st.text_input("Name", key="Name")
    Phone = st.text_input("Phone", key="Phone")
    Email = st.text_input("Email", key="Email")
with c2:
    City = st.text_input("City", key="City")
    State = st.text_input("State (2-letter)", key="State")

st.subheader("3. Objective")
c3a, c3b = st.columns(2)
with c3a:
    application_type = st.radio("Are you seeking a job or apprenticeship?", ["Apprenticeship","Job"], horizontal=True, index=0)
    trade = st.selectbox("What trade are you aiming for?", TRADE_TAXONOMY, index=TRADE_TAXONOMY.index("Electrician – Inside (01)"), key="SelectedTrade")
with c3b:
    wk_pitch = st.text_input("10-second pitch (what you want them to know):", st.session_state.get("Pitch",""))
    st.session_state["Pitch"] = wk_pitch

wk_objective_final = st.text_area("Objective (1–2 sentences — prefilled; edit if you want):", st.session_state.get("Objective_Final",""))

st.subheader("4. Skills (auto suggestions + editable)")
suggested_skills = suggest_transferable_skills_from_text(combined_text)
quick_transfer = st.multiselect("Quick Add: transferable skills from your uploads", SKILL_CANON, default=suggested_skills)
Skills_Transferable = st.text_area("Transferable Skills (comma/newline):", st.session_state.get("Skills_Transferable",""))
Skills_JobSpecific  = st.text_area("Job-Specific Skills (comma/newline):", st.session_state.get("Skills_JobSpecific",""))
Skills_SelfManagement = st.text_area("Self-Management Skills (comma/newline):", st.session_state.get("Skills_SelfManagement",""))

st.subheader("5. Work Experience – Job 1")
J1c = st.text_input("Job 1 – Company:", key="Job1_Company")
J1cs = st.text_input("Job 1 – City/State:", key="Job1_CityState")
J1d = st.text_input("Job 1 – Dates (e.g., 2023-06 – Present):", key="Job1_Dates")
J1t = st.text_input("Job 1 – Title:", key="Job1_Title")
seed1=[]
for key_role, arr in ROLE_TO_CONSTR_BULLETS.items():
    if key_role in (st.session_state.get("Job1_Title","").lower()):
        seed1 = arr[:MAX_BULLETS_PER_JOB]
J1du = st.text_area("Job 1 – Duties/Accomplishments (1–4 bullets):", key="Job1_Duties", value=st.session_state.get("Job1_Duties","\n".join(seed1)), height=120)

st.subheader("5. Work Experience – Job 2")
J2c = st.text_input("Job 2 – Company:", key="Job2_Company")
J2cs = st.text_input("Job 2 – City/State:", key="Job2_CityState")
J2d = st.text_input("Job 2 – Dates:", key="Job2_Dates")
J2t = st.text_input("Job 2 – Title:", key="Job2_Title")
seed2=[]
for key_role, arr in ROLE_TO_CONSTR_BULLETS.items():
    if key_role in (st.session_state.get("Job2_Title","").lower()):
        seed2 = arr[:MAX_BULLETS_PER_JOB]
J2du = st.text_area("Job 2 – Duties/Accomplishments (1–4 bullets):", key="Job2_Duties", value=st.session_state.get("Job2_Duties","\n".join(seed2)), height=120)

st.subheader("5. Work Experience – Job 3")
J3c = st.text_input("Job 3 – Company:", key="Job3_Company")
J3cs = st.text_input("Job 3 – City/State:", key="Job3_CityState")
J3d = st.text_input("Job 3 – Dates:", key="Job3_Dates")
J3t = st.text_input("Job 3 – Title:", key="Job3_Title")
seed3=[]
for key_role, arr in ROLE_TO_CONSTR_BULLETS.items():
    if key_role in (st.session_state.get("Job3_Title","").lower()):
        seed3 = arr[:MAX_BULLETS_PER_JOB]
J3du = st.text_area("Job 3 – Duties/Accomplishments (1–4 bullets):", key="Job3_Duties", value=st.session_state.get("Job3_Duties","\n".join(seed3)), height=120)

st.subheader("6. Certifications")
Certifications = st.text_area(
    "List certifications (comma/newline). If none, write 'None yet' or what you plan to get.",
    st.session_state.get("Certifications","OSHA-10, Flagger (WA), Forklift, CPR")
)

st.subheader("7. Education")
st.write("Reverse order. Include city/state, dates, and credential/diploma.")
E1s = st.text_input("School/Program 1:", key="Edu1_School"); E1cs = st.text_input("City/State 1:", key="Edu1_CityState")
E1d = st.text_input("Dates 1:", key="Edu1_Dates"); E1c = st.text_input("Certificate/Diploma 1:", key="Edu1_Credential")
E2s = st.text_input("School/Program 2:", key="Edu2_School"); E2cs = st.text_input("City/State 2:", key="Edu2_CityState")
E2d = st.text_input("Dates 2:", key="Edu2_Dates"); E2c = st.text_input("Certificate/Diploma 2:", key="Edu2_Credential")

st.markdown("**Final Checklist**")
st.markdown("""
- [ ] One page only  
- [ ] Professional font (10–12 pt)  
- [ ] Saved as PDF  
- [ ] Reviewed by peer  
- [ ] Reviewed by instructor  
""")

# ─────────────────────────────────────────────────────────
# Cover Letter (optional)
# ─────────────────────────────────────────────────────────
st.markdown("---")
st.subheader("Cover Letter (optional)")
CL_Company = st.text_input("Company/Employer:","")
CL_Role    = st.text_input("Role Title:", f"{st.session_state.get('SelectedTrade','Apprenticeship')} apprentice")
CL_Location= st.text_input("Company Location (City, State):","")
CL_Highlights = st.text_area("Optional: bullet highlights (comma/newline/• allowed):","Reliable • Safety-focused • Coachable")

# ─────────────────────────────────────────────────────────
# Generate Docs
# ─────────────────────────────────────────────────────────
if st.button("Generate Resume + Cover Letter + Instructor Packet", type="primary"):
    problems=[]
    if not (st.session_state.get("Name","").strip()):
        problems.append("Name is required.")
    if not (st.session_state.get("Phone","").strip() or st.session_state.get("Email","").strip()):
        problems.append("At least one contact method (Phone or Email) is required.")
    if not tpl_bytes:
        problems.append("Resume template missing. Upload resume_app_template.docx in the sidebar or keep it at repo root.")
    if problems:
        st.error(" | ".join(problems))
        st.stop()

    # Merge Quick Add into Transferable
    skills_transfer_final = st.session_state.get("Skills_Transferable","")
    if quick_transfer:
        skills_transfer_final = (skills_transfer_final + (", " if skills_transfer_final.strip() else "") + ", ".join(quick_transfer))

    trade = st.session_state.get("SelectedTrade","Electrician – Inside (01)")
    form = {
        "Name": st.session_state.get("Name",""), "City": st.session_state.get("City",""), "State": st.session_state.get("State",""),
        "Phone": st.session_state.get("Phone",""), "Email": st.session_state.get("Email",""),
        "Pitch": st.session_state.get("Pitch",""),
        "Objective_Final": st.session_state.get("Objective_Final", wk_objective_final) or wk_objective_final,
        "Skills_Transferable": skills_transfer_final,
        "Skills_JobSpecific": st.session_state.get("Skills_JobSpecific",""),
        "Skills_SelfManagement": st.session_state.get("Skills_SelfManagement",""),
        "Certifications": st.session_state.get("Certifications", Certifications),
        "Other_Work": st.session_state.get("Other_Work",""), "Volunteer": st.session_state.get("Volunteer",""),
    }
    # Jobs
    for i in [1,2,3]:
        form[f"Job{i}_Company"]=st.session_state.get(f"Job{i}_Company","")
        form[f"Job{i}_CityState"]=st.session_state.get(f"Job{i}_CityState","")
        form[f"Job{i}_Dates"]=st.session_state.get(f"Job{i}_Dates","")
        form[f"Job{i}_Title"]=st.session_state.get(f"Job{i}_Title","")
        form[f"Job{i}_Duties"]=st.session_state.get(f"Job{i}_Duties","")
    # Education
    for i in [1,2]:
        form[f"Edu{i}_School"]=st.session_state.get(f"Edu{i}_School","")
        form[f"Edu{i}_CityState"]=st.session_state.get(f"Edu{i}_CityState","")
        form[f"Edu{i}_Dates"]=st.session_state.get(f"Edu{i}_Dates","")
        form[f"Edu{i}_Credential"]=st.session_state.get(f"Edu{i}_Credential","")

    # Resume
    try:
        resume_ctx = build_resume_context(form, trade)
        resume_bytes = render_docx_with_template(tpl_bytes, resume_ctx)
    except Exception as e:
        st.error(f"Resume template rendering failed: {e}")
        st.stop()

    # Cover Letter
    cover_bytes = build_cover_letter_docx({
        "name": form["Name"], "city": form["City"], "state": form["State"], "phone": clean_phone(form["Phone"]), "email": clean_email(form["Email"]),
        "company": CL_Company, "role": CL_Role, "location": CL_Location,
        "trade_label": trade, "strength": CL_Highlights,
        "application_type": "Apprenticeship" if "Apprentice" in CL_Role.title() else "Job",
    })

    # Instructor Packet (Workshop reflections + full text of docs + TOC + Sources + Roadmap slice)
    reflections = {
        "Three construction-resume items (vs traditional)": wk_q1,
    }
    # Merge *all* uploads into packet
    url_fetch_files = []
    for nb in url_fetches:
        # Include the fetched file in packet with a friendly name
        url_fetch_files.append(NamedBytesIO(nb.getvalue() if hasattr(nb, "getvalue") else nb.read(), getattr(nb, "name", "downloaded.txt")))
    merged_docs_for_packet = list(pathway_uploads or []) + list(prev_resume_files or []) + list(jd_files or []) + url_fetch_files
    packet_bytes = build_pathway_packet_docx({"name": form["Name"]}, trade, "Apprenticeship" if application_type=="Apprenticeship" else "Job", merged_docs_for_packet, reflections)

    safe_name = (form["Name"] or "Student").replace(" ","_")
    st.download_button("Download Resume (DOCX)", data=resume_bytes,
                       file_name=f"{safe_name}_Resume.docx",
                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                       use_container_width=True)
    st.download_button("Download Cover Letter (DOCX)", data=cover_bytes,
                       file_name=f"{safe_name}_Cover_Letter.docx",
                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                       use_container_width=True)
    st.download_button("Download Instructor Pathway Packet (DOCX)", data=packet_bytes,
                       file_name=f"{safe_name}_Instructor_Pathway_Packet.docx",
                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                       use_container_width=True)

    # Intake CSV — fixed column order
    csv_fields = [
        "Name","City","State","Phone","Email",
        "Objective_Final",
        "Skills_Transferable","Skills_JobSpecific","Skills_SelfManagement",
        "Certifications",
        "Job1_Company","Job1_CityState","Job1_Dates","Job1_Title","Job1_Duties",
        "Job2_Company","Job2_CityState","Job2_Dates","Job2_Title","Job2_Duties",
        "Job3_Company","Job3_CityState","Job3_Dates","Job3_Title","Job3_Duties",
        "Edu1_School","Edu1_CityState","Edu1_Dates","Edu1_Credential",
        "Edu2_School","Edu2_CityState","Edu2_Dates","Edu2_Credential"
    ]
    buf=io.StringIO(); w=csv.writer(buf)
    w.writerow(csv_fields); w.writerow([form.get(k,"") for k in csv_fields])
    st.download_button("Download Intake CSV", data=buf.getvalue().encode("utf-8"),
                       file_name=f"{safe_name}_Workshop_Intake.csv", mime="text/csv",
                       use_container_width=True)
    st.success("Generated. Autofill: header + jobs + skills + objective. Packet includes TOC, Sources, and Roadmap slice when provided.")
